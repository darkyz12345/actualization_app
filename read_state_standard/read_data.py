import time
from docx import Document
import re
from .utils import clean_text
from PyQt5.QtCore import pyqtSignal
from collections.abc import Iterator
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter


def read_data_from_docx(source_path: str, progress_bar: pyqtSignal) -> list[str]:
    """
    Reads state standard names from a .docx file.

    :param source_path: Path to the source .docx file containing the state standard names.
    :return: A list of state standard names read from the source file.
    """
    data = {'O`zDSt': r'(O`z DSt \S+|O`zDSt \S+)',
            'ГОСТ': r'(ГОСТ ISO \S+|ГОСТ Р МЭК \S+|ГОСТ Р \S+|ГОСТ IEC \S+|ГОСТ МЭК \S+|ГОСТ EN \S+|ГОСТ \S+|ГОСТ\S+)',
            'ISO': r'(ISO \S+)',
            'UzTR.': r'(UzTR.\S+)'}
    word_doc = Document(source_path)
    table = word_doc.tables[-1]
    length = len(table.rows[1:])
    standard_set = set()
    for i, row in enumerate(table.rows[1:], start=1):
        standard = row.cells[3].text
        if standard:
            standard = clean_text(standard)
            for pattern in data.values():
                matches = re.findall(pattern, standard)
                standard_set.update(matches)
        progress_percent = int(i / length * 100)
        progress_bar.emit(progress_percent)
    # with open(destination_path, 'w', encoding='utf-8') as f:
    #     for s in standard_set:
    #         f.write(f'{s}\n')
    return list(standard_set)

def save_to_txt(standards: Iterator[str], destination: str, length: int, progress_bar: pyqtSignal):
    standards_iter = enumerate(standards, start=1)
    if '.txt' in destination:
        with open(destination, 'w', encoding='utf-8') as file:
            for i, standard in standards_iter:
                file.write(f'{i}) {standard}\n')
                time.sleep(0.0005)
                progress_percent = int(i / length * 100)
                progress_bar.emit(progress_percent)
        return
    elif '.xlsx' in destination:
        center_align = Alignment(horizontal='center', vertical='center')
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '№'
        ws['B1'] = 'Найденные ГОСТЫ'
        ws.title = 'ГОСТы'
        for i, standard in standards_iter:
            ws.append([i, standard])
            time.sleep(0.0005)
            progress_percent = int(i / length * 100)
            progress_bar.emit(progress_percent)
        for row in ws['A']:
            row.alignment = center_align
        for row in ws['B']:
            row.alignment = center_align
        for col in ws.columns:
            max_length = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = max_length + 2  # +2 для отступа
        wb.save(destination)
        return
    document = Document()
    document.add_paragraph('Найденные ГОСТы').alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(1, 2, "Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = 'Название ГОСТа'
    for cell in hdr_cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, standard in standards_iter:
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = standard
        time.sleep(0.0005)
        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        progress_percent = int(i / length * 100)
        progress_bar.emit(progress_percent)
    document.save(destination)
    return



def read_standards_from_txt(source_path: str) -> list[str]:
    """
    Read state standards from .txt file
    :param source_path: Path to the source .txt file containing the state standard names.
    :return: A list of state standard names read from the source file.
    """
    data = []
    with open(source_path, 'rt', encoding='utf-8') as file:
        data = [f.strip().replace(';', '') for f in file]
    return data

def read_data_for_search(source_path: str, progress_bar: pyqtSignal) -> list[str]:
    fake_progress = 0
    # print(source_path)
    res = []
    if 'txt' in source_path:
        # print('Find')
        with open(source_path, 'rt', encoding='utf-8') as file:
            for line in file:
                res.append(line.strip())
                time.sleep(0.0005)
                progress_bar.emit(fake_progress)
                if fake_progress < 90:
                    fake_progress += 0.005
    elif 'docx' in source_path:
        doc = Document(source_path)
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    continue
                time.sleep(0.0005)
                progress_bar.emit(fake_progress)
                if fake_progress < 90:
                    fake_progress += 0.5
                res.append(row.cells[1].text.strip())
    else:
        wb = load_workbook(source_path)
        ws = wb.active
        for row in ws.iter_rows(min_row=2):
            time.sleep(0.0005)
            progress_bar.emit(fake_progress)
            if fake_progress < 90:
                fake_progress += 0.5
            value = row[1].value
            res.append(str(value).strip())
        progress_bar.emit(100)
    return res


if __name__ == '__main__':
    src_path = '/media/stepan/Disk/python_projects/2025/actualization_app/standard.docx'
    dest_path = '/media/stepan/Disk/python_projects/2025/actualization_app/standards.txt'
    # read_data_from_docx(src_path, dest_path)
    # print(read_standards_from_txt(dest_path))
