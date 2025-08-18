from docx import Document
from PyQt5 import QtCore
from PyQt5.QtGui import QStandardItemModel
from openpyxl import Workbook
from openpyxl.utils import get_column_letter

def save_table_to_file(filename: str, model: QStandardItemModel, progress_bar: QtCore.pyqtSignal):
    if filename.endswith('.docx'):
        print("Headers:", [model.headerData(i, QtCore.Qt.Horizontal, QtCore.Qt.DisplayRole)
                           for i in range(model.columnCount())])
        doc = Document()
        total_rows = model.rowCount()

        # создаём таблицу: +1 строка под заголовки
        table = doc.add_table(rows=total_rows + 1, cols=model.columnCount())
        table.style = 'Table Grid'

        # === шапка таблицы ===
        for col in range(model.columnCount()):
            hdr = model.headerData(col, QtCore.Qt.Horizontal, QtCore.Qt.DisplayRole)
            cell = table.cell(0, col)
            cell.text = str(hdr) if hdr else ""
            # делаем заголовки жирными
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        # === данные ===
        for row in range(total_rows):
            for col in range(model.columnCount()):
                item = model.item(row, col)
                table.cell(row + 1, col).text = item.text() if item else ""

            # обновление прогресса
            progress_bar.emit(int((row + 1) * 100 / total_rows))

        doc.save(filename)

    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Данные"

        total_rows = model.rowCount()

        # === шапка таблицы ===
        headers = [model.headerData(col, QtCore.Qt.Horizontal, QtCore.Qt.DisplayRole)
                   for col in range(model.columnCount())]
        ws.append(headers)

        # === данные ===
        for row in range(total_rows):
            row_data = []
            for col in range(model.columnCount()):
                item = model.item(row, col)
                row_data.append(item.text() if item else "")
            ws.append(row_data)

            # обновляем прогресс
            progress_bar.emit(int((row + 1) * 100 / total_rows))

        # === автоширина колонок ===
        for col in ws.columns:
            max_length = 0
            col_letter = get_column_letter(col[0].column)  # буква колонки (A, B, C...)
            for cell in col:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            # +2 для небольшого отступа
            ws.column_dimensions[col_letter].width = max_length + 2

        # сохраняем Excel
        wb.save(filename)